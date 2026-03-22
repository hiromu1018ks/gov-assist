"""Generate .docx from corrected text using python-docx (§6.2)."""
import re
import io

from docx import Document

# Bullet patterns (§6.2): three distinct patterns, each matched independently
# ・ (katakana middle dot): optional space after (Japanese convention: no space)
# - (hyphen): requires space after (to avoid matching "--" or "-text")
# N. (number + dot): requires space after
_BULLET_RE = re.compile(
    r"^(?:(?P<katakana>・)\s*|(?P<hyphen>-)\s+|(?P<number>\d+)\.\s+)(?P<text>.+)$"
)

# python-docx built-in list style names
_LIST_BULLET_STYLE = "List Bullet"
_LIST_NUMBER_STYLE = "List Number"


def _add_paragraph(doc, text: str) -> None:
    """Add a single paragraph to the document, detecting bullet/numbered style."""
    match = _BULLET_RE.match(text)
    if match:
        bullet_text = match.group("text")
        if match.group("number") is not None:
            doc.add_paragraph(bullet_text, style=_LIST_NUMBER_STYLE)
        else:
            # Both katakana dot and hyphen use List Bullet style
            doc.add_paragraph(bullet_text, style=_LIST_BULLET_STYLE)
    else:
        doc.add_paragraph(text)


def generate_docx(corrected_text: str, document_type: str) -> bytes:  # noqa: ARG001
    """Generate .docx bytes from corrected text.

    Rules (§6.2):
    - Empty lines (\\n\\n) → paragraph separators
    - Single newlines within a block → kept in the same paragraph
    - Lines starting with "・", "-", or "数字+." → list style applied
    - Bullet marker is stripped from the paragraph text (the style adds its own marker)
    - Plain text base, no original formatting reproduction

    Args:
        corrected_text: The proofread/corrected text content.
        document_type: Document type identifier (email/report/official/other).
                      Currently unused but kept for future per-type styling.

    Returns:
        .docx file content as bytes.

    Raises:
        ValueError: If corrected_text is empty or whitespace-only.
    """
    if corrected_text is None or not corrected_text.strip():
        raise ValueError("corrected_text must not be empty")

    doc = Document()

    # Split on empty lines (paragraph boundaries)
    # A blank line = two consecutive newlines = paragraph separator
    blocks = re.split(r"\n\s*\n", corrected_text.strip())

    for block in blocks:
        # Each block becomes one paragraph.
        # Single \\n within a block is kept as-is in the paragraph text.
        _add_paragraph(doc, block.strip())

    # Write to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
