# backend/tests/test_docx_exporter.py
"""Tests for services/docx_exporter.py — .docx generation service (§6.2)."""
import io

import pytest
from docx import Document

from services.docx_exporter import generate_docx


class TestGenerateDocxBasic:
    def test_returns_bytes(self):
        result = generate_docx("テスト", "official")
        assert isinstance(result, bytes)

    def test_valid_docx_structure(self):
        result = generate_docx("テスト", "official")
        # ZIP header (PK\x03\x04) indicates valid docx/zip
        assert result[:4] == b"PK\x03\x04"

    def test_single_paragraph(self):
        result = generate_docx("一つの段落です。", "official")
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "一つの段落です。"

    def test_empty_string_raises(self):
        with pytest.raises(ValueError):
            generate_docx("", "official")

    def test_whitespace_only_raises(self):
        with pytest.raises(ValueError):
            generate_docx("   \n\n  ", "official")


class TestParagraphSplitting:
    def test_empty_line_creates_new_paragraph(self):
        text = "第一段落\n\n第二段落"
        result = generate_docx(text, "official")
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) == 2
        assert doc.paragraphs[0].text == "第一段落"
        assert doc.paragraphs[1].text == "第二段落"

    def test_multiple_empty_lines_collapsed(self):
        text = "第一段落\n\n\n\n第二段落"
        result = generate_docx(text, "official")
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) == 2

    def test_single_newline_does_not_split(self):
        text = "一行目\n二行目"
        result = generate_docx(text, "official")
        doc = Document(io.BytesIO(result))
        # Single \n within a paragraph should keep it as one paragraph
        assert len(doc.paragraphs) == 1

    def test_trailing_empty_lines_ignored(self):
        text = "本文\n\n\n"
        result = generate_docx(text, "official")
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "本文"


class TestBulletDetection:
    def test_katakana_middle_dot(self):
        text = "・項目1\n\n・項目2"
        result = generate_docx(text, "official")
        doc = Document(io.BytesIO(result))
        assert doc.paragraphs[0].style.name.startswith("List")
        assert doc.paragraphs[1].style.name.startswith("List")

    def test_hyphen_bullet(self):
        text = "- 項目A\n\n- 項目B"
        result = generate_docx(text, "official")
        doc = Document(io.BytesIO(result))
        assert doc.paragraphs[0].style.name.startswith("List")

    def test_numbered_list(self):
        text = "1. 第一項\n\n2. 第二項\n\n3. 第三項"
        result = generate_docx(text, "official")
        doc = Document(io.BytesIO(result))
        for p in doc.paragraphs:
            assert p.style.name.startswith("List")

    def test_mixed_content(self):
        text = "見出し\n\n・箇条書き1\n\n・箇条書き2\n\n本文"
        result = generate_docx(text, "official")
        doc = Document(io.BytesIO(result))
        assert doc.paragraphs[0].text == "見出し"
        assert not doc.paragraphs[0].style.name.startswith("List")
        assert doc.paragraphs[1].style.name.startswith("List")
        assert doc.paragraphs[2].style.name.startswith("List")
        assert doc.paragraphs[3].text == "本文"
        assert not doc.paragraphs[3].style.name.startswith("List")

    def test_bullet_stripped_from_text(self):
        """箇条書き記号は python-docx の List Bullet スタイルが付与するため、元テキストからは除去する"""
        text = "・項目テキスト"
        result = generate_docx(text, "official")
        doc = Document(io.BytesIO(result))
        assert doc.paragraphs[0].text == "項目テキスト"

    def test_numbered_list_stripped(self):
        text = "1. 項目テキスト"
        result = generate_docx(text, "official")
        doc = Document(io.BytesIO(result))
        assert doc.paragraphs[0].text == "項目テキスト"

    def test_double_hyphen_not_bullet(self):
        """「--」は箇条書きとして扱わない（「-」単体のみが箇条書き記号）"""
        text = "-- テスト"
        result = generate_docx(text, "official")
        doc = Document(io.BytesIO(result))
        assert not doc.paragraphs[0].style.name.startswith("List")

    def test_digit_without_dot_not_numbered(self):
        """「数字のみ」（ドットなし）は箇条書きとして扱わない"""
        text = "123 テスト"
        result = generate_docx(text, "official")
        doc = Document(io.BytesIO(result))
        assert not doc.paragraphs[0].style.name.startswith("List")


class TestDocumentTypeHandling:
    def test_all_document_types(self):
        for doc_type in ["email", "report", "official", "other"]:
            result = generate_docx("テスト", doc_type)
            assert isinstance(result, bytes)
            assert result[:4] == b"PK\x03\x04"
